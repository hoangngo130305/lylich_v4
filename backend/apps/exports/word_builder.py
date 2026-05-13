"""
Build a Mẫu 2-KNĐ Word document from a Profile instance.
Uses python-docx. Returns BytesIO ready for HttpResponse.
"""
import io
import traceback

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.utils.timezone import localdate


# ── Helpers ───────────────────────────────────────────────────────────────────

def _safe(val, default=''):
    if val is None:
        return default
    return str(val)


def _bold(run):
    run.bold = True
    return run


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    return p


def _section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return p


def _field(doc, label, value, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run_label = p.add_run(f'{label}: ')
    run_label.bold = True
    p.add_run(_safe(value))
    return p


def _table_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = _safe(text)
    return row


def _period_str(from_month, from_year, to_month, to_year, is_present=False):
    def fmt(m, y):
        if y and m:
            return f'{int(m):02d}/{y}'
        return str(y) if y else ''
    start = fmt(from_month, from_year)
    end = 'nay' if is_present else fmt(to_month, to_year)
    if start and end:
        return f'{start} – {end}'
    if start:
        return f'{start} – …'
    return ''


def _fmt_date(d):
    if not d:
        return ''
    return d.strftime('%d/%m/%Y')


# ── Main builder ──────────────────────────────────────────────────────────────

def build_profile_docx(profile, sections=None) -> io.BytesIO:
    """
    Build a Mẫu 2-KNĐ Word document.

    sections: list of section keys to include, e.g. ['basic', 'family', 'work'].
              None → include all sections.
    """
    print(f"[EXPORT] generating word for profile {profile.id} – {profile.full_name}")

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(2)

    _heading(doc, 'SƠ YẾU LÝ LỊCH ĐẢNG VIÊN', level=1)
    _heading(doc, '(Mẫu 2-KNĐ)', level=2)
    doc.add_paragraph()

    def _include(key):
        return sections is None or key in sections

    # ── Section A: Thông tin cá nhân ─────────────────────────────────────────
    if _include('basic'):
        _section_title(doc, 'A. THÔNG TIN CÁ NHÂN')

        _field(doc, 'Họ và tên', profile.full_name)
        _field(doc, 'Tên gọi khác', profile.full_name_other)
        _field(doc, 'Ngày sinh', _fmt_date(profile.dob))
        _field(doc, 'Giới tính', profile.get_gender_display() if profile.gender else '')

        # Dân tộc — FK first, fall back to free-text
        ethnic = ''
        if profile.ethnic_group_id:
            try:
                ethnic = profile.ethnic_group.name
            except Exception:
                pass
        ethnic = ethnic or _safe(profile.ethnic_group_other)
        _field(doc, 'Dân tộc', ethnic)

        # Tôn giáo
        religion = ''
        if profile.religion_id:
            try:
                religion = profile.religion.name
            except Exception:
                pass
        religion = religion or _safe(profile.religion_other)
        _field(doc, 'Tôn giáo', religion)

        # CCCD / phone / email come from the related User record
        user = None
        try:
            user = profile.user
        except Exception:
            print(f"[EXPORT] Warning: could not access profile.user\n{traceback.format_exc()}")

        _field(doc, 'CCCD/CMND', getattr(user, 'cccd', '') or '')
        _field(doc, 'Số điện thoại', getattr(user, 'phone', '') or '')
        _field(doc, 'Email', getattr(user, 'email', '') or '')

        _field(doc, 'Quê quán', profile.hometown_detail)
        _field(doc, 'Nơi sinh', profile.birth_place_detail)
        _field(doc, 'Nơi ở hiện tại', profile.current_address)
        doc.add_paragraph()

    # ── Section B: Trình độ học vấn ──────────────────────────────────────────
    if _include('education'):
        _section_title(doc, 'B. TRÌNH ĐỘ HỌC VẤN')

        edu_name = ''
        if profile.edu_level_id:
            try:
                edu_name = profile.edu_level.name
            except Exception:
                pass
        edu_name = edu_name or _safe(profile.general_edu_level)
        _field(doc, 'Trình độ học vấn', edu_name)

        _field(doc, 'Chuyên ngành', profile.edu_major)
        _field(doc, 'Trường', profile.edu_school)

        pol_name = ''
        if profile.political_level_id:
            try:
                pol_name = profile.political_level.name
            except Exception:
                pass
        pol_name = pol_name or _safe(profile.political_level_detail)
        _field(doc, 'Trình độ lý luận chính trị', pol_name)

        foreign_lang = profile.foreign_languages or profile.foreign_language_name or ''
        _field(doc, 'Ngoại ngữ', foreign_lang)
        _field(doc, 'Tin học', profile.it_level)
        doc.add_paragraph()

    # ── Section C: Quá trình công tác ────────────────────────────────────────
    if _include('work'):
        _section_title(doc, 'C. QUÁ TRÌNH CÔNG TÁC')
        try:
            work_qs = profile.work_history.order_by('from_year', 'from_month', 'sort_order')
            if work_qs.exists():
                tbl = doc.add_table(rows=1, cols=4)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr = tbl.rows[0].cells
                hdr[0].text = 'Từ – Đến'
                hdr[1].text = 'Chức vụ'
                hdr[2].text = 'Đơn vị'
                hdr[3].text = 'Ghi chú'
                for w in work_qs:
                    period = w.period_text or _period_str(
                        w.from_month, w.from_year, w.to_month, w.to_year, w.is_present
                    )
                    _table_row(tbl, [period, w.job_title, w.employer, w.notes])
        except Exception:
            print(f"[EXPORT] Warning: work_history load failed\n{traceback.format_exc()}")
        doc.add_paragraph()

    # ── Section D: Gia đình ───────────────────────────────────────────────────
    if _include('family'):
        _section_title(doc, 'D. THÀNH PHẦN GIA ĐÌNH')
        try:
            family_qs = profile.family_members.filter(deleted_at__isnull=True).order_by('sort_order', 'relationship')
            if family_qs.exists():
                tbl = doc.add_table(rows=1, cols=5)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr = tbl.rows[0].cells
                hdr[0].text = 'Quan hệ'
                hdr[1].text = 'Họ tên'
                hdr[2].text = 'Năm sinh'
                hdr[3].text = 'Nghề nghiệp'
                hdr[4].text = 'Nơi ở'
                for fm in family_qs:
                    birth_year = fm.dob.year if fm.dob else (fm.birth_year or '')
                    _table_row(tbl, [
                        fm.get_relationship_display(),
                        fm.full_name,
                        birth_year,
                        fm.occupation,
                        fm.current_address,
                    ])
        except Exception:
            print(f"[EXPORT] Warning: family_members load failed\n{traceback.format_exc()}")
        doc.add_paragraph()

    # ── Section E: Khen thưởng & Kỷ luật ────────────────────────────────────
    if _include('awards'):
        _section_title(doc, 'E. KHEN THƯỞNG & KỶ LUẬT')
        try:
            awards_qs = profile.awards.order_by('issued_year', 'issued_month', 'sort_order')
            if awards_qs.exists():
                tbl = doc.add_table(rows=1, cols=4)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr = tbl.rows[0].cells
                hdr[0].text = 'Năm'
                hdr[1].text = 'Loại'
                hdr[2].text = 'Nội dung'
                hdr[3].text = 'Cấp'
                for a in awards_qs:
                    _table_row(tbl, [
                        a.issued_year,
                        a.get_type_display(),
                        a.content,
                        a.level,
                    ])
        except Exception:
            print(f"[EXPORT] Warning: awards load failed\n{traceback.format_exc()}")
        doc.add_paragraph()

    # ── Section F: Tự nhận xét ───────────────────────────────────────────────
    if _include('self_assessment'):
        _section_title(doc, 'F. TỰ NHẬN XÉT – ĐÁNH GIÁ')
        p = doc.add_paragraph(_safe(profile.self_assessment_text))
        p.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    # ── Signature block ───────────────────────────────────────────────────────
    today = localdate()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'Ngày {today.day} tháng {today.month} năm {today.year}')
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _bold(p2.add_run('NGƯỜI KHAI KÝ TÊN'))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    print(f"[EXPORT] Word document built OK for profile {profile.id}")
    return buf
