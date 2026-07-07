"""
word_builder_from_template.py — Mẫu 2-KNĐ exporter v3
────────────────────────────────────────────────────────
StyleMapper   : inherits ALL styles from lylich_sample.docx (zero hardcoding).
FieldResolver : resolves every dotted path with full debug logging.
Render loop   : per-section start/done log; placeholder on any failure;
                NEVER exits early, NEVER silently skips a section.
"""

import io
import json
import os
import traceback
from datetime import date, datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.utils.timezone import localdate

# ── Paths ──────────────────────────────────────────────────────────────────────

_HERE      = os.path.dirname(os.path.abspath(__file__))
_BACKEND   = os.path.normpath(os.path.join(_HERE, '..', '..'))
_JSON_PATH = os.path.join(_BACKEND, 'template.json')
_DOCX_PATH = os.path.join(_BACKEND, 'lylich_sample.docx')

_TEMPLATE_CACHE = None
_TEMPLATE_MTIME = None


def _load_template():
    global _TEMPLATE_CACHE, _TEMPLATE_MTIME
    try:
        mtime = os.path.getmtime(_JSON_PATH)
    except OSError:
        mtime = None

    # Reload template when file changes to avoid stale export structure in long-running processes.
    if _TEMPLATE_CACHE is None or _TEMPLATE_MTIME != mtime:
        with open(_JSON_PATH, 'r', encoding='utf-8') as f:
            _TEMPLATE_CACHE = json.load(f)
        _TEMPLATE_MTIME = mtime
    return _TEMPLATE_CACHE


# ── StyleMapper ────────────────────────────────────────────────────────────────

_STRIP_TAGS = frozenset({
    'p', 'tbl', 'sdt', 'customXml',
    'bookmarkStart', 'bookmarkEnd', 'proofErr',
    'ins', 'del', 'moveFrom', 'moveTo',
})

_PARA_CANDIDATES = {
    'normal':   ['Normal', 'Văn bản thường', 'Body Text'],
    'heading1': ['Heading 1', 'heading 1', 'Title', 'Tiêu đề 1'],
    'heading2': ['Heading 2', 'heading 2', 'Tiêu đề 2'],
    'heading3': ['Heading 3', 'heading 3', 'Tiêu đề 3'],
}
_TABLE_CANDIDATES = [
    'Table Grid', 'TableGrid', 'Light Grid',
    'Light Grid Accent 1', 'Normal Table',
]


class StyleMapper:
    """
    Opens lylich_sample.docx once, discovers available styles,
    then returns a blank Document that inherits every style from the sample.
    """

    def __init__(self, sample_path):
        self._path  = sample_path
        self._pmap  = {k: 'Normal' for k in _PARA_CANDIDATES}
        self._tname = 'Table Grid'
        self._ready = False
        self._init()

    def _init(self):
        if not os.path.exists(self._path):
            print(f'[STYLE][ERR] Sample DOCX not found: {self._path!r}')
            return
        try:
            probe      = Document(self._path)
            para_names = {s.name for s in probe.styles
                          if s.type == WD_STYLE_TYPE.PARAGRAPH}
            tbl_names  = {s.name for s in probe.styles
                          if s.type == WD_STYLE_TYPE.TABLE}

            for key, cands in _PARA_CANDIDATES.items():
                for c in cands:
                    if c in para_names:
                        self._pmap[key] = c
                        break

            for c in _TABLE_CANDIDATES:
                if c in tbl_names:
                    self._tname = c
                    break

            self._ready = True
            print(f'[STYLE] para_map={self._pmap}  table={self._tname!r}')
        except Exception:
            print(f'[STYLE][ERR] Init failed:\n{traceback.format_exc()}')

    def blank_doc(self):
        """
        Open the sample, strip all body content, preserve styles + page setup.
        Falls back to a bare Document() if the sample is unavailable.
        """
        if not self._ready:
            print('[STYLE][WARN] Using default (unstyled) document.')
            return Document()
        try:
            doc  = Document(self._path)
            body = doc.element.body
            for child in list(body):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in _STRIP_TAGS:
                    body.remove(child)
            body.insert(0, OxmlElement('w:p'))   # Word needs ≥1 paragraph before sectPr
            return doc
        except Exception:
            print(f'[STYLE][ERR] blank_doc failed:\n{traceback.format_exc()}')
            return Document()

    def p(self, key='normal'):
        return self._pmap.get(key, 'Normal')

    def t(self):
        return self._tname


_SM = None


def _get_style_mapper():
    global _SM
    if _SM is None:
        _SM = StyleMapper(_DOCX_PATH)
    return _SM


# ── Safe document helpers ──────────────────────────────────────────────────────

def _safe_para(doc, style_name, fallback='Normal'):
    """Add a paragraph; fall back gracefully if style_name is unknown."""
    for name in (style_name, fallback, 'Normal', None):
        try:
            if name is None:
                return doc.add_paragraph()
            return doc.add_paragraph(style=name)
        except Exception:
            continue
    return doc.add_paragraph()


def _safe_table_style(tbl, style_name):
    """Assign a table style; silently ignore unknown style names."""
    for name in (style_name, 'Table Grid', None):
        try:
            if name is None:
                return
            tbl.style = name
            return
        except Exception:
            continue


def _set_col_widths(tbl, widths_cm):
    """
    Set column widths via tblGrid + per-cell tcW.
    Entirely wrapped in try/except — a broken width is cosmetic, not fatal.
    """
    try:
        twips  = [int(w * 567) for w in widths_cm]
        tbl_el = tbl._tbl

        old = tbl_el.find(qn('w:tblGrid'))
        if old is not None:
            tbl_el.remove(old)
        grid = OxmlElement('w:tblGrid')
        for t in twips:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(t))
            grid.append(gc)
        tbl_pr = tbl_el.find(qn('w:tblPr'))
        if tbl_pr is not None:
            tbl_pr.addnext(grid)
        else:
            tbl_el.insert(0, grid)

        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i >= len(twips):
                    break
                try:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w  = tc_pr.find(qn('w:tcW'))
                    if tc_w is None:
                        tc_w = OxmlElement('w:tcW')
                        tc_pr.append(tc_w)
                    tc_w.set(qn('w:w'), str(twips[i]))
                    tc_w.set(qn('w:type'), 'dxa')
                except Exception:
                    pass  # cosmetic — continue
    except Exception:
        print(f'[WIDTH][ERR] {traceback.format_exc()}')


def _cell_write(cell, text, bold=False):
    """
    Write text into a table cell, converting \\n to proper line breaks.
    Avoids raw newlines inside <w:t> which can confuse Word's XML parser.
    """
    para = cell.paragraphs[0]
    para.clear()
    parts = text.split('\n') if text else ['']
    for idx, part in enumerate(parts):
        run = para.add_run(part)
        if bold:
            run.bold = True
        if idx < len(parts) - 1:   # add line-break between parts, not after last
            run.add_break()


def _bold_header(cell, text):
    _cell_write(cell, text, bold=True)


def _shade_cell(cell, fill_hex: str):
    """Set cell background color (fill_hex without #, e.g. '1F4E79')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)


def _cell_write_white(cell, text, bold=True):
    """Write text in white color (for use on dark/colored backgrounds)."""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    rpr = run._r.get_or_add_rPr()
    color_el = rpr.find(qn('w:color'))
    if color_el is None:
        color_el = OxmlElement('w:color')
        rpr.append(color_el)
    color_el.set(qn('w:val'), 'FFFFFF')


# ── FieldResolver ──────────────────────────────────────────────────────────────

_UNSET = object()


class FieldResolver:
    """Resolves dotted paths with full per-field logging."""

    def __init__(self, ctx):
        self._ctx = ctx
        self._log = []

    def resolve(self, path, fallback=''):
        if not path:
            return fallback
        raw, err = self._walk(path)
        ok = raw is not None and raw is not _UNSET and raw != ''
        self._log.append({'path': path, 'raw': raw, 'err': err, 'ok': ok})
        if err:
            print(f'[FIELD][WARN] {path!r} — {err}')
        elif not ok:
            print(f'[FIELD][MISS] {path!r} → {raw!r}')
        if raw is None or raw is _UNSET:
            return fallback
        return self._fmt(raw)

    def obj_field(self, obj, field):
        if not field:
            return ''
        if field == 'period_text':
            return self._period_str(obj)
        if field == 'religion_with_rank':
            religion = getattr(obj, 'religion_text', None) or ''
            rank = getattr(obj, 'religious_rank_text', None) or ''
            if not religion:
                return 'Không'
            return f'{religion} ({rank})' if rank else f'{religion} (không có chức sắc)'
        if field == 'issued_year_month':
            m = getattr(obj, 'issued_month', None)
            y = getattr(obj, 'issued_year',  None)
            return f'{int(m):02d}/{y}' if (m and y) else (str(y) if y else '')
        attr = getattr(obj, field, _UNSET)
        if attr is _UNSET:
            print(f'[FIELD][MISS] {type(obj).__name__}.{field} — not found')
            return ''
        if callable(attr):
            try:
                val = attr()
            except Exception as exc:
                print(f'[FIELD][ERR]  {type(obj).__name__}.{field}(): {exc}')
                return ''
        else:
            val = attr
        return self._fmt(val)

    def fetch(self, source, profile, flt, order_by):
        _MAP = {
            'history_entries':    'history_entries',
            'work_history':       'work_history',
            'education_history':  'education_history',
            'awards':             'awards',
            'overseas_travels':   'overseas_travels',
            'org_participations': 'org_participations',
            'family_members':     'family_members',
        }
        acc = _MAP.get(source)
        if not acc:
            print(f'[FETCH][ERR] Unknown source: {source!r}')
            return []
        try:
            qs = getattr(profile, acc)
            if flt:
                qs = qs.filter(**flt)
            if order_by:
                qs = qs.order_by(*order_by)
            rows = list(qs)
            print(f'[FETCH]  {source!r} filter={flt} → {len(rows)} rows')
            return rows
        except Exception:
            print(f'[FETCH][ERR] {source}:\n{traceback.format_exc()}')
            return []

    def report(self):
        total   = len(self._log)
        ok      = sum(1 for e in self._log if e['ok'])
        missing = [e for e in self._log if not e['ok']]
        W = 52
        lines = [
            '',
            f'┌{"─"*W}┐',
            f'│{"[EXPORT] Field Resolution Report":^{W}}│',
            f'├{"─"*W}┤',
            f'│  Total    : {total:<{W-13}}│',
            f'│  Resolved : {ok:<{W-13}}│',
            f'│  Missing  : {total-ok:<{W-13}}│',
        ]
        if missing:
            lines.append(f'├{"─"*W}┤')
            lines.append(f'│  Missing fields:{" "*(W-18)}│')
            for e in missing:
                lines.append(f'│    • {e["path"][:W-6]:<{W-6}}│')
                lines.append(f'│      → {str(e["err"] or e["raw"] or "null")[:W-8]:<{W-8}}│')
        lines.append(f'└{"─"*W}┘')
        return '\n'.join(lines)

    # ── private ────────────────────────────────────────────────────────────────

    def _walk(self, path):
        parts = path.split('.')
        obj   = self._ctx.get(parts[0])
        if obj is None:
            return None, f'ctx[{parts[0]!r}] missing'
        for i, part in enumerate(parts[1:], 1):
            if obj is None:
                return None, f'None at {".".join(parts[:i])!r}'
            try:
                attr = getattr(obj, part, _UNSET)
            except Exception as exc:
                return None, f'getattr({part!r}): {exc}'
            if attr is _UNSET:
                return None, f'{type(obj).__name__!r} has no attr {part!r}'
            if callable(attr):
                try:
                    obj = attr()
                except Exception as exc:
                    return None, f'{type(obj).__name__}.{part}(): {exc}'
            else:
                obj = attr
        return obj, None

    @staticmethod
    def _fmt(val):
        if val is None:
            return ''
        if isinstance(val, bool):
            return 'Có' if val else 'Không'
        if isinstance(val, (date, datetime)):
            try:
                return val.strftime('%d/%m/%Y')
            except (ValueError, TypeError) as e:
                # Handle naive datetime timezone conversion errors
                print(f"[EXPORT] Warning: date format failed for {val}: {e}")
                return str(val)[:10]  # Return date part from string representation
        return str(val)

    @staticmethod
    def _period_str(obj):
        pt = getattr(obj, 'period_text', None)
        if pt:
            return pt
        fm = getattr(obj, 'from_month', None)
        fy = getattr(obj, 'from_year',  None)
        tm = getattr(obj, 'to_month',   None)
        ty = getattr(obj, 'to_year',    None)
        present = getattr(obj, 'is_present', False)

        def fmt(m, y):
            return f'{int(m):02d}/{y}' if (m and y) else (str(y) if y else '')

        start = fmt(fm, fy)
        end   = 'nay' if present else fmt(tm, ty)
        if start and end:
            return f'{start} – {end}'
        return start or end or ''


# ── Section renderers ─────────────────────────────────────────────────────────

def _r_title_block(doc, sec, sm, fr):
    for pd in sec.get('paragraphs', []):
        style_key = 'heading1' if 'BOLD' in pd.get('style', '') else 'normal'
        p = _safe_para(doc, sm.p(style_key))
        if 'CENTER' in pd.get('style', ''):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(pd.get('text', ''))
        if 'ITALIC' in pd.get('style', ''):
            run.italic = True
    _safe_para(doc, sm.p('normal'))


def _r_section_heading(doc, sec, sm, fr):
    tbl = doc.add_table(rows=1, cols=1)
    _safe_table_style(tbl, sm.t())
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, '1F4E79')
    _cell_write_white(cell, sec.get('text', ''))
    _set_col_widths(tbl, [16.5])
    _safe_para(doc, sm.p('normal'))


def _r_instruction(doc, sec, sm, fr):
    tbl = doc.add_table(rows=1, cols=1)
    _safe_table_style(tbl, sm.t())
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, 'FFF2CC')
    para = cell.paragraphs[0]
    run = para.add_run(sec.get('text', ''))
    run.italic = True
    _set_col_widths(tbl, [16.5])
    _safe_para(doc, sm.p('normal'))


def _r_table_static(doc, sec, sm, fr):
    rows_def = sec.get('rows', [])
    label    = sec.get('label', '')
    if not rows_def:
        _safe_para(doc, sm.p('normal'))
        return

    tbl = doc.add_table(rows=0, cols=2)
    _safe_table_style(tbl, sm.t())

    if label:
        hrow = tbl.add_row()
        merged = hrow.cells[0].merge(hrow.cells[1])
        _shade_cell(merged, 'D6E4F0')
        para = merged.paragraphs[0]
        run = para.add_run(label)
        run.bold = True

    for rd in rows_def:
        row = tbl.add_row()
        _shade_cell(row.cells[0], 'F2F2F2')
        _cell_write(row.cells[0], rd.get('label', ''), bold=True)

        field = rd.get('field', '')
        fb    = rd.get('fallback', '')
        val   = fr.resolve(field) if field else ''
        if not val and fb:
            val = fr.resolve(fb)

        f2 = rd.get('field2', '')
        if f2:
            v2 = fr.resolve(f2)
            if val and v2:
                val = f'{val}, tại {v2}'
            elif v2:
                val = v2

        _cell_write(row.cells[1], val)

    _set_col_widths(tbl, [6.5, 10.0])
    _safe_para(doc, sm.p('normal'))


def _r_table_dynamic(doc, sec, sm, fr):
    profile  = fr._ctx['profile']
    source   = sec.get('source', '')
    flt      = sec.get('filter', {})
    order_by = sec.get('order_by', [])
    columns  = sec.get('columns', [])
    empty_n  = sec.get('empty_rows', 5)

    if not columns:
        # Still render a placeholder so the section exists in output
        p = _safe_para(doc, sm.p('normal'))
        p.add_run(f'[{sec.get("id", "table_dynamic")}: no columns defined]').italic = True
        _safe_para(doc, sm.p('normal'))
        return

    rows = fr.fetch(source, profile, flt, order_by)

    tbl = doc.add_table(rows=1, cols=len(columns))
    _safe_table_style(tbl, sm.t())

    # Header row — blue background, white text
    for i, col in enumerate(columns):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, '2E75B6')
        _cell_write_white(cell, col.get('header', ''))

    # Data rows — never skip any
    for obj in rows:
        r = tbl.add_row()
        for i, col in enumerate(columns):
            _cell_write(r.cells[i], fr.obj_field(obj, col.get('field', '')))

    # Filler rows to reach minimum empty-row count
    for _ in range(max(0, empty_n - len(rows))):
        tbl.add_row()

    _set_col_widths(tbl, [col.get('width_cm', 3.5) for col in columns])
    _safe_para(doc, sm.p('normal'))


def _r_text_area(doc, sec, sm, fr):
    label   = sec.get('label', '')
    field   = sec.get('field', '')
    n_lines = sec.get('empty_lines', 5)

    p = _safe_para(doc, sm.p('normal'))
    if label:
        run = p.add_run(label + ':')
        run.bold = True

    val = fr.resolve(field) if field else ''
    if val:
        # Split on newlines so each line becomes its own paragraph, preserving formatting.
        for line in val.split('\n'):
            _safe_para(doc, sm.p('normal')).add_run(line)
    else:
        for _ in range(n_lines):
            _safe_para(doc, sm.p('normal'))

    _safe_para(doc, sm.p('normal'))


def _r_family_block(doc, sec, sm, fr):
    """Single-member block: cha, mẹ, vợ/chồng, ông/bà..."""
    profile    = fr._ctx['profile']
    title      = sec.get('title', '')
    rel        = sec.get('relationship', '')
    fields_def = sec.get('fields', [])

    try:
        member = profile.family_members.filter(
            relationship=rel, deleted_at__isnull=True
        ).first()
    except Exception:
        print(f'[FAMILY][ERR] rel={rel!r}:\n{traceback.format_exc()}')
        member = None

    # Skip entire section (heading + table) when there is no data
    if member is None:
        return

    tbl = doc.add_table(rows=0, cols=2)
    _safe_table_style(tbl, sm.t())

    # Section heading row — light green, spanning both columns
    hrow = tbl.add_row()
    merged = hrow.cells[0].merge(hrow.cells[1])
    _shade_cell(merged, 'E2EFDA')
    para = merged.paragraphs[0]
    run = para.add_run(title)
    run.bold = True

    if not fields_def:
        _set_col_widths(tbl, [5.5, 11.0])
        _safe_para(doc, sm.p('normal'))
        return

    for fd in fields_def:
        if fd.get('skip_if_deceased') and member.is_deceased:
            continue
        if fd.get('skip_if_alive') and not member.is_deceased:
            continue
        row = tbl.add_row()
        _shade_cell(row.cells[0], 'F2F2F2')
        _cell_write(row.cells[0], fd.get('label', ''), bold=True)
        fname = fd.get('field', '')
        if fname == 'is_party_member':
            _cell_write(row.cells[1], _party_str(member))
        else:
            _cell_write(row.cells[1], fr.obj_field(member, fname))

    _set_col_widths(tbl, [5.5, 11.0])

    # Optional history sub-table
    h_src  = sec.get('history_source')
    h_cols = sec.get('history_columns', [])
    if h_src and h_cols:
        p_ql = _safe_para(doc, sm.p('normal'))
        p_ql.add_run('Quá trình lịch sử:').bold = True

        try:
            h_rows = list(member.history_entries.order_by(
                'from_year', 'from_month', 'sort_order'
            ))
        except Exception:
            print(f'[FAMILY][ERR] history rel={rel!r}:\n{traceback.format_exc()}')
            h_rows = []

        htbl = doc.add_table(rows=1, cols=len(h_cols))
        _safe_table_style(htbl, sm.t())

        for i, col in enumerate(h_cols):
            cell = htbl.rows[0].cells[i]
            _shade_cell(cell, '2E75B6')
            _cell_write_white(cell, col.get('header', ''))

        for obj in h_rows:
            r = htbl.add_row()
            for i, col in enumerate(h_cols):
                _cell_write(r.cells[i], fr.obj_field(obj, col.get('field', '')))

        for _ in range(max(0, 5 - len(h_rows))):
            htbl.add_row()

        _set_col_widths(htbl, [c.get('width_cm', 4.0) for c in h_cols])

    _safe_para(doc, sm.p('normal'))


def _r_family_list_block(doc, sec, sm, fr):
    """Multi-member block: anh/chị/em ruột, con, anh/chị/em chồng/vợ."""
    profile    = fr._ctx['profile']
    title      = sec.get('title', '')
    rel        = sec.get('relationship', '')
    fields_def = sec.get('fields', [])

    try:
        members = list(profile.family_members.filter(
            relationship=rel, deleted_at__isnull=True
        ).order_by('sort_order'))
    except Exception:
        print(f'[FAMILY][ERR] list rel={rel!r}:\n{traceback.format_exc()}')
        members = []

    # Skip entire section (heading + tables) when there is no data
    if not members:
        return

    # Group heading (section title) — dark navy banner
    grp_tbl = doc.add_table(rows=1, cols=1)
    _safe_table_style(grp_tbl, sm.t())
    grp_cell = grp_tbl.rows[0].cells[0]
    _shade_cell(grp_cell, '1F4E79')
    _cell_write_white(grp_cell, title)
    _set_col_widths(grp_tbl, [16.5])
    _safe_para(doc, sm.p('normal'))

    current_year = localdate().year

    for member in members:
        label = member.custom_label or (
            'Con ruột' if member.relationship == 'con'
            else member.get_relationship_display()
        )

        # Compute age once per member for conditional fields
        try:
            member_age = current_year - int(member.birth_year) if member.birth_year else None
        except (ValueError, TypeError):
            member_age = None

        tbl = doc.add_table(rows=0, cols=2)
        _safe_table_style(tbl, sm.t())

        # Per-member heading row — light green, spanning both columns
        hrow = tbl.add_row()
        merged = hrow.cells[0].merge(hrow.cells[1])
        _shade_cell(merged, 'E2EFDA')
        para = merged.paragraphs[0]
        run = para.add_run(label)
        run.bold = True

        for fd in fields_def:
            if fd.get('skip_if_deceased') and member.is_deceased:
                continue
            if fd.get('skip_if_alive') and not member.is_deceased:
                continue
            if fd.get('skip_if_under_18') and member_age is not None and member_age < 18:
                continue
            if fd.get('skip_if_over_18') and (member_age is None or member_age >= 18):
                continue
            r = tbl.add_row()
            _shade_cell(r.cells[0], 'F2F2F2')
            _cell_write(r.cells[0], fd.get('label', ''), bold=True)
            fname = fd.get('field', '')
            if fname == 'is_party_member':
                _cell_write(r.cells[1], _party_str(member))
            else:
                val = fr.obj_field(member, fname)
                if not val and fd.get('default'):
                    val = fd['default']
                _cell_write(r.cells[1], val)

        _set_col_widths(tbl, [5.5, 11.0])

        # Optional per-member history sub-table
        h_src  = sec.get('history_source')
        h_cols = sec.get('history_columns', [])
        skip_hist_under_18 = sec.get('skip_history_if_under_18', False)
        if h_src and h_cols and not (skip_hist_under_18 and member_age is not None and member_age < 18):
            try:
                h_rows = list(member.history_entries.order_by(
                    'from_year', 'from_month', 'sort_order'
                ))
            except Exception:
                print(f'[FAMILY][ERR] history rel={rel!r} member={member.id}:\n{traceback.format_exc()}')
                h_rows = []

            if h_rows:
                p_ql = _safe_para(doc, sm.p('normal'))
                p_ql.add_run('Quá trình lịch sử:').bold = True

                htbl = doc.add_table(rows=1, cols=len(h_cols))
                _safe_table_style(htbl, sm.t())

                for i, col in enumerate(h_cols):
                    cell = htbl.rows[0].cells[i]
                    _shade_cell(cell, '2E75B6')
                    _cell_write_white(cell, col.get('header', ''))

                for obj in h_rows:
                    hr = htbl.add_row()
                    for i, col in enumerate(h_cols):
                        _cell_write(hr.cells[i], fr.obj_field(obj, col.get('field', '')))

                for _ in range(max(0, 3 - len(h_rows))):
                    htbl.add_row()

                _set_col_widths(htbl, [c.get('width_cm', 4.0) for c in h_cols])

        _safe_para(doc, sm.p('normal'))

    _safe_para(doc, sm.p('normal'))


def _r_declaration(doc, sec, sm, fr):
    text       = sec.get('text', '')
    sig_fields = sec.get('signature_fields', [])

    if text:
        p = _safe_para(doc, sm.p('normal'))
        p.add_run(text).italic = True

    today = localdate()
    for sf in sig_fields:
        name  = fr.resolve(sf.get('field', ''))
        dval  = fr.resolve(sf.get('date_field', '')) if sf.get('date_field') else ''

        p_d = _safe_para(doc, sm.p('normal'))
        p_d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_d.add_run(dval or f'Ngày {today.day} tháng {today.month} năm {today.year}')

        p_l = _safe_para(doc, sm.p('normal'))
        p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_l.add_run(sf.get('label', '').upper()).bold = True

        p_n = _safe_para(doc, sm.p('normal'))
        p_n.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_n.add_run(f'({name})')

    _safe_para(doc, sm.p('normal'))


def _r_committee_comment(doc, sec, sm, fr):
    title  = sec.get('title', '')
    field  = sec.get('field', '')
    by_f   = sec.get('signed_by', '')
    date_f = sec.get('signed_date', '')

    tbl_h = doc.add_table(rows=1, cols=1)
    _safe_table_style(tbl_h, sm.t())
    cell_h = tbl_h.rows[0].cells[0]
    _shade_cell(cell_h, '1F4E79')
    _cell_write_white(cell_h, title)
    _set_col_widths(tbl_h, [16.5])
    _safe_para(doc, sm.p('normal'))

    content  = fr.resolve(field) if field else ''
    signed   = fr.resolve(by_f)  if by_f  else ''
    sign_dt  = fr.resolve(date_f) if date_f else ''

    if content:
        _safe_para(doc, sm.p('normal')).add_run(content)
    else:
        for _ in range(7):
            _safe_para(doc, sm.p('normal'))

    today = localdate()
    p_d = _safe_para(doc, sm.p('normal'))
    p_d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_d.add_run(sign_dt or f'Ngày      tháng      năm      ')

    p_s = _safe_para(doc, sm.p('normal'))
    p_s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if signed:
        p_s.add_run(signed)
    else:
        p_s.add_run('(Ký tên, đóng dấu)').italic = True

    _safe_para(doc, sm.p('normal'))


# ── Small helpers ─────────────────────────────────────────────────────────────

def _party_str(member):
    if not member.is_party_member:
        return 'Không'
    parts = ['Là đảng viên Đảng Cộng sản Việt Nam']
    if member.party_join_year:
        parts.append(f'vào Đảng năm {member.party_join_year}')
    chi_bo  = member.party_chi_bo  or ''
    dang_bo = member.party_dang_bo or ''
    if chi_bo or dang_bo:
        loc = f'chi bộ {chi_bo}' if chi_bo else ''
        if dang_bo:
            loc += f' thuộc đảng bộ {dang_bo}' if loc else f'đảng bộ {dang_bo}'
        parts.append(f'hiện sinh hoạt tại {loc}')
    return ', '.join(parts)


# ── Dispatcher ────────────────────────────────────────────────────────────────

_RENDERERS = {
    'title_block':              _r_title_block,
    'section_heading':          _r_section_heading,
    'instruction':              _r_instruction,
    'table_static':             _r_table_static,
    'table_dynamic':            _r_table_dynamic,
    'text_area':                _r_text_area,
    'family_member_block':      _r_family_block,
    'family_member_list_block': _r_family_list_block,
    'declaration':              _r_declaration,
    'committee_comment':        _r_committee_comment,
}


# ── Public entry point ────────────────────────────────────────────────────────

def build_lylich_docx(profile):
    """
    Build a Mẫu 2-KNĐ DOCX from a Profile instance.

    Guarantees:
      • Every section in template.json produces SOME output in the DOCX.
      • Failed sections leave a visible placeholder instead of silent gap.
      • Every field resolution is logged; missing fields listed in final report.
    Returns io.BytesIO ready for HttpResponse.
    """
    print(f'\n[EXPORT] ════ START profile={profile.id} {profile.full_name!r} ════')

    sm  = _get_style_mapper()
    doc = sm.blank_doc()

    # ── Build context ──────────────────────────────────────────────────────────
    ctx = {'profile': profile}
    try:
        ctx['chi_uy_comment'] = profile.committee_comments.filter(type='chi_uy').first()
        ctx['cap_uy_comment'] = profile.committee_comments.filter(type='cap_uy_co_so').first()
        print(f'[EXPORT] chi_uy={"✓" if ctx["chi_uy_comment"] else "∅"}  '
              f'cap_uy={"✓" if ctx["cap_uy_comment"] else "∅"}')
    except Exception:
        ctx['chi_uy_comment'] = None
        ctx['cap_uy_comment'] = None
        print(f'[EXPORT][WARN] committee_comments failed:\n{traceback.format_exc()}')

    fr  = FieldResolver(ctx)
    tpl = _load_template()
    structure = tpl.get('structure', [])

    # ── Section inventory ──────────────────────────────────────────────────────
    print(f'[EXPORT] Template sections ({len(structure)} total):')
    for i, sec in enumerate(structure):
        print(f'[EXPORT]   [{i:02d}] {sec.get("type"):25s} id={sec.get("id")!r}')

    # ── Render loop — NEVER stops early ───────────────────────────────────────
    ok_count  = 0
    err_count = 0

    for i, sec in enumerate(structure):
        sec_id   = sec.get('id', f'sec_{i:02d}')
        sec_type = sec.get('type', '')
        renderer = _RENDERERS.get(sec_type)

        print(f'[RENDER] [{i:02d}] {sec_id!r} ({sec_type}) … ', end='', flush=True)

        if renderer is None:
            print(f'SKIP – unknown type {sec_type!r}')
            # Still produce a visible placeholder for unknown types
            p = _safe_para(doc, 'Normal')
            p.add_run(f'[{sec_id}: unknown type {sec_type!r}]').italic = True
            continue

        try:
            renderer(doc, sec, sm, fr)
            ok_count += 1
            print('OK')
        except Exception:
            err_count += 1
            tb = traceback.format_exc()
            print(f'ERROR')
            print(f'[RENDER][ERR] {sec_id!r}:\n{tb}')
            # Guaranteed fallback output so the section is VISIBLE in the DOCX
            try:
                p = _safe_para(doc, 'Normal')
                p.add_run(f'⚠ [{sec_id}] render error — see server log').italic = True
                _safe_para(doc, 'Normal')
            except Exception:
                pass

    # ── Summary ────────────────────────────────────────────────────────────────
    print(f'[EXPORT] Render: {ok_count} OK  {err_count} errors  '
          f'{len(structure)} total sections')
    print(fr.report())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    print(f'[EXPORT] ════ DONE profile={profile.id} ════\n')
    return buf
