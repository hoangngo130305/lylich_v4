"""
convert_docx_to_template.py
────────────────────────────
Reads lylich_sample.docx, walks the body in DOM order,
and outputs template.json + prints a structured log.

Run:
    python convert_docx_to_template.py
"""

import json
import re
import os
import sys

# Force UTF-8 stdout so Vietnamese chars don't crash on Windows cp1252
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml.ns import qn

# ── Config ────────────────────────────────────────────────────────────────────
_HERE = os.path.dirname(os.path.abspath(__file__))

# Find the DOCX regardless of double-extension quirk
def _find_docx():
    candidates = [
        os.path.join(_HERE, "lylich_sample.docx"),
        os.path.join(_HERE, "lylich_sample.docx.docx"),
        os.path.join(_HERE, "backend", "lylich_sample.docx"),
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    # Fallback: scan _HERE and backend/
    for search_dir in (_HERE, os.path.join(_HERE, "backend")):
        if not os.path.isdir(search_dir):
            continue
        for f in os.listdir(search_dir):
            if f.lower().endswith(".docx") and "sample" in f.lower():
                return os.path.join(search_dir, f)
    return None

DOCX_PATH   = _find_docx()
OUTPUT_JSON = os.path.join(_HERE, "template.json")

# ── Helpers ───────────────────────────────────────────────────────────────────

def get_cell_text(cell):
    """Return full text of a cell (may contain multiple paragraphs)."""
    return "\n".join(p.text for p in cell.paragraphs if p.text.strip())


def get_run_bold(paragraph):
    """True if ALL non-empty runs are bold."""
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.bold for r in runs)


def classify_style(paragraph):
    """Return a simplified style tag."""
    name = paragraph.style.name.lower()
    if "heading 1" in name or "title" in name:
        return "HEADING1"
    if "heading 2" in name:
        return "HEADING2"
    if "heading 3" in name:
        return "HEADING3"
    if get_run_bold(paragraph):
        return "BOLD"
    return "NORMAL"


def detect_placeholders(text):
    """Return list of {{…}} placeholders found in text."""
    return re.findall(r"\{\{[^}]+\}\}", text)


# ── Placeholder → Django model mapping ───────────────────────────────────────
PLACEHOLDER_MAP = {
    # Mẫu 2-KNĐ field labels → Django profile/user fields
    "họ và tên":           "{{profile.full_name}}",
    "ho va ten":           "{{profile.full_name}}",
    "full_name":           "{{profile.full_name}}",
    "tên gọi khác":        "{{profile.full_name_other}}",
    "ngày sinh":           "{{profile.dob}}",
    "nơi sinh":            "{{profile.birth_place_detail}}",
    "giới tính":           "{{profile.get_gender_display}}",
    "dân tộc":             "{{profile.ethnic_group.name}}",
    "tôn giáo":            "{{profile.religion.name}}",
    "quê quán":            "{{profile.hometown_detail}}",
    "nơi ở hiện nay":      "{{profile.current_address}}",
    "nơi ở hiện tại":      "{{profile.current_address}}",
    "địa chỉ":             "{{profile.current_address}}",
    "cccd":                "{{profile.user.cccd}}",
    "cmnd":                "{{profile.user.cccd}}",
    "số điện thoại":       "{{profile.user.phone}}",
    "điện thoại":          "{{profile.user.phone}}",
    "email":               "{{profile.user.email}}",
    "trình độ học vấn":    "{{profile.edu_level.name}}",
    "trình độ văn hóa":    "{{profile.general_edu_level}}",
    "chuyên ngành":        "{{profile.edu_major}}",
    "trường":              "{{profile.edu_school}}",
    "trình độ lý luận":    "{{profile.political_level.name}}",
    "ngoại ngữ":           "{{profile.foreign_languages}}",
    "tin học":             "{{profile.it_level}}",
    "nghề nghiệp":         "{{profile.occupation}}",
    "chức vụ":             "{{profile.job_title}}",
    "nơi công tác":        "{{profile.workplace}}",
    "ngày vào đoàn":       "{{profile.youth_union_date}}",
    "kết nạp đảng":        "{{profile.rejoin_first_party_join_date}}",
    "điểm ai":             "{{profile.ai_score}}",
    "số hồ sơ":            "{{profile.profile_number}}",
    "trạng thái":          "{{profile.get_status_display}}",
    "ngày nộp":            "{{profile.submitted_at}}",
    "cán bộ phụ trách":    "{{profile.officer_in_charge.full_name}}",
}

DYNAMIC_REGIONS = {
    # keywords that signal a dynamic table (list from DB)
    "quá trình công tác":  "work_history",
    "quá trình học tập":   "education_history",
    "thành phần gia đình": "family_members",
    "quan hệ gia đình":    "family_members",
    "khen thưởng":         "awards",
    "kỷ luật":             "awards",
    "ra nước ngoài":       "overseas_travels",
    "người thân ở nước ngoài": "overseas_relatives",
    "tổ chức":             "org_participations",
}

DYNAMIC_COLUMN_MAP = {
    "work_history": {
        "từ":       "period_from",
        "đến":      "period_to",
        "thời gian": "period_text",
        "chức vụ":  "job_title",
        "đơn vị":   "employer",
        "nơi làm":  "employer",
        "ghi chú":  "notes",
    },
    "education_history": {
        "từ":       "from_year",
        "đến":      "to_year",
        "thời gian": "period_text",
        "trường":   "school",
        "ngành":    "major",
        "bằng":     "certificate",
        "ghi chú":  "notes",
    },
    "family_members": {
        "quan hệ":   "get_relationship_display",
        "họ tên":    "full_name",
        "năm sinh":  "birth_year",
        "nghề nghiệp": "occupation",
        "nơi ở":     "current_address",
        "đảng viên": "is_party_member",
        "chính trị": "political_status",
        "ghi chú":   "notes",
    },
    "awards": {
        "năm":      "issued_year",
        "loại":     "get_type_display",
        "nội dung": "content",
        "cấp":      "level",
        "cơ quan":  "issuer",
    },
    "overseas_travels": {
        "thời gian": "period_text",
        "nước":      "country",
        "mục đích":  "purpose",
        "cơ quan":   "sponsoring_org",
    },
}

# ── Main parser ───────────────────────────────────────────────────────────────

def guess_source(header_rows):
    """Try to detect what DB table a table maps to."""
    combined = " ".join(
        cell.lower()
        for row in header_rows
        for cell in row
    )
    for keyword, source in DYNAMIC_REGIONS.items():
        if keyword in combined:
            return source
    return None


def map_columns(source, header_row):
    """Map header cell labels to model field names."""
    col_map = DYNAMIC_COLUMN_MAP.get(source, {})
    mapping = []
    for cell in header_row:
        key = cell.lower().strip()
        field = col_map.get(key, key.replace(" ", "_"))
        mapping.append([cell, field])
    return mapping


def label_to_placeholder(label_text):
    """Map a label string to a Django template placeholder."""
    key = label_text.lower().strip().rstrip(":：")
    for k, v in PLACEHOLDER_MAP.items():
        if k in key or key in k:
            return v
    return None


def parse_table(tbl, t_index, preceding_heading=None):
    """
    Parse a docx table into either:
      - 'table_static'  : key-value or mixed static layout
      - 'table_dynamic' : repeating data rows (list from DB)
    Returns a dict section entry.
    """
    all_rows = []
    for row in tbl.rows:
        cells = [get_cell_text(c) for c in row.cells]
        # de-duplicate merged cells (docx repeats merged cell text)
        deduped = []
        prev = None
        for c in cells:
            if c != prev:
                deduped.append(c)
            prev = c
        all_rows.append(deduped)

    if not all_rows:
        return None

    # Detect if this is a dynamic data table
    heading_text = (preceding_heading or "").lower()
    source = guess_source(all_rows[:2])
    if source is None:
        source = guess_source([[heading_text]])

    if source:
        header_row = all_rows[0] if all_rows else []
        mapping    = map_columns(source, header_row)
        section = {
            "type":    "table_dynamic",
            "source":  source,
            "header":  header_row,
            "mapping": mapping,
            "raw_rows": all_rows,
        }
        print(f"  → DYNAMIC table[{t_index}] source={source!r} cols={header_row}")
    else:
        # Static table — detect placeholders in value cells
        rows_with_placeholders = []
        for row in all_rows:
            annotated = []
            for i, cell in enumerate(row):
                ph = label_to_placeholder(cell) if i == 0 and len(row) > 1 else None
                annotated.append({"text": cell, "placeholder": ph})
            rows_with_placeholders.append(annotated)

        section = {
            "type": "table_static",
            "rows": rows_with_placeholders,
            "raw_rows": all_rows,
        }
        print(f"  → STATIC table[{t_index}] {len(all_rows)} rows x {max(len(r) for r in all_rows)} cols")

    return section


def parse_document(docx_path):
    doc = Document(docx_path)
    body = doc.element.body

    structure   = []
    placeholders = {}
    p_index      = 0
    t_index      = 0
    last_heading = None

    print(f"\n{'='*60}")
    print(f"  Parsing: {docx_path}")
    print(f"  Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}")
    print(f"{'='*60}\n")

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        # ── Paragraph ──────────────────────────────────────────────
        if tag == "p":
            if p_index < len(doc.paragraphs):
                para  = doc.paragraphs[p_index]
                text  = para.text.strip()
                style = classify_style(para)

                if text:
                    section = {
                        "type":  "paragraph",
                        "style": style,
                        "text":  text,
                    }

                    # Collect any {{placeholder}} already in doc
                    for ph in detect_placeholders(text):
                        placeholders[ph] = ph

                    # Track last heading for table context
                    if style in ("HEADING1", "HEADING2", "HEADING3", "BOLD"):
                        last_heading = text

                    structure.append(section)
                    print(f"[P:{p_index:03d}] [{style:8s}] {text[:100]}")

            p_index += 1

        # ── Table ──────────────────────────────────────────────────
        elif tag == "tbl":
            if t_index < len(doc.tables):
                tbl     = doc.tables[t_index]
                section = parse_table(tbl, t_index, preceding_heading=last_heading)
                if section:
                    structure.append(section)
            t_index += 1

    # ── Auto-collect placeholders from static table labels ─────────
    for sec in structure:
        if sec.get("type") == "table_static":
            for row in sec["rows"]:
                for cell in row:
                    if cell.get("placeholder"):
                        label = cell["text"].rstrip(":：").strip()
                        placeholders[label] = cell["placeholder"]

    return structure, placeholders


def build_template_json(structure, placeholders):
    """Build the final template dict — strip raw_rows from output."""
    clean_sections = []
    for sec in structure:
        s = {k: v for k, v in sec.items() if k != "raw_rows"}
        clean_sections.append(s)

    return {
        "template_name": "ly_lich_mau2_knd",
        "version":       "1.0",
        "description":   "Sơ yếu lý lịch Đảng viên – Mẫu 2-KNĐ",
        "structure":     clean_sections,
        "placeholders":  placeholders,
        "django_models": {
            "profile":           "apps.profiles.models.Profile",
            "user":              "apps.accounts.models.User",
            "family_members":    "apps.family.models.FamilyMember",
            "work_history":      "apps.timelines.models.WorkHistory",
            "education_history": "apps.timelines.models.EducationHistory",
            "awards":            "apps.timelines.models.Award",
            "overseas_travels":  "apps.timelines.models.OverseasTravel",
            "org_participations":"apps.timelines.models.OrgParticipation",
        },
    }


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if DOCX_PATH is None:
        print("ERROR: lylich_sample.docx not found. Files in dir:")
        for f in os.listdir(_HERE):
            print(" ", f)
        sys.exit(1)

    structure, placeholders = parse_document(DOCX_PATH)
    template = build_template_json(structure, placeholders)

    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump(template, f, ensure_ascii=False, indent=2)

    print(f"\n{'='*60}")
    print(f"  Sections parsed : {len(structure)}")
    print(f"  Placeholders    : {len(placeholders)}")
    print(f"  Output          : {OUTPUT_JSON}")
    print(f"{'='*60}")
